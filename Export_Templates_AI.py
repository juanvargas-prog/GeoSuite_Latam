import pandas as pd
import geopandas as gpd
import gspread
from google.oauth2.service_account import Credentials
import os
from datetime import datetime
import certifi
import json
from dotenv import load_dotenv

def exportar_gdf(bd, nombre_capa, gdf, formato, base_folder="C:/export_latam"):
    """
    Crea una carpeta base si no existe, dentro crea una subcarpeta con el nombre_variable + fecha,
    y exporta el GeoDataFrame en shapefile o CSV según el formato indicado.
    
    Parámetros:
        nombre_variable (str): nombre dinámico para la carpeta
        gdf (GeoDataFrame): datos a exportar
        formato (str): 'shp' o 'csv'
        base_folder (str): carpeta raíz donde se guardará todo
    """
    # 1. Crear carpeta base si no existe
    os.makedirs(base_folder, exist_ok=True)

    # 2. Crear subcarpeta con nombre_variable + fecha
    fecha = datetime.now().strftime("%Y%m%d")
    subfolder = os.path.join(base_folder, f"{bd}_{fecha}")
    os.makedirs(subfolder, exist_ok=True)

    # 3. Exportar según formato
    if formato.lower() == "shp":
        file_final = f"{nombre_capa}.{formato}"
        gdf.to_file(os.path.join(subfolder, file_final))
    elif formato.lower() == "csv":
        file_final = f"{nombre_capa}.{formato}"
        gdf.drop(columns="geom").to_csv(os.path.join(subfolder, file_final), index=False)
    elif formato.lower() == "gpkg":
        file_final = f"{nombre_capa}.gpkg"
        gdf.to_file(os.path.join(subfolder, file_final), driver="GPKG", layer=nombre_capa)
    else:
        raise ValueError("Formato no soportado. Usa 'shp', 'csv' o 'gpkg'.")

    print(f"Archivo exportado en: {subfolder}")

def main_exportar_capas(bd, valor, pares, columna_seleccionada, engine, formato):
    """
    Exporta capas desde la BD usando la estructura del Google Sheet.

    pares: lista de tuplas (nombre_capa_bd, nombre_estructura_sheet)
           Cada par indica qué capa del servidor exportar y qué fila del Sheet
           usar para obtener las columnas con TRUE en la columna de producción.
    """

    # Forzamos a las librerías a buscar los certificados donde certifi dice que están
    os.environ['REQUESTS_CA_BUNDLE'] = certifi.where()
    os.environ['SSL_CERT_FILE'] = certifi.where()

    # Configuración de las variables desde el archivo local .env
    ruta_plugin = os.path.dirname(__file__)
    ruta_env = os.path.join(ruta_plugin, ".env")
    load_dotenv(dotenv_path=ruta_env)

    # 2. Leer la variable
    service_account_str = os.getenv("SERVICE_ACCOUNT_INFO")

    if service_account_str:
        SERVICE_ACCOUNT_INFO = json.loads(service_account_str)
        print("Credenciales cargadas correctamente")
    else:
        print("No se encontro la variable de entorno SERVICE_ACCOUNT_INFO")        
    
    # 1. CONFIGURACIÓN DE GOOGLE SHEETS
    scope = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
    creds = Credentials.from_service_account_info(SERVICE_ACCOUNT_INFO, scopes=scope)
    client = gspread.authorize(creds)

    # Abrir el documento por nombre o por URL
    spreadsheet = client.open("07_Estructura Cartografía Base")
    sheet = spreadsheet.worksheet(valor)

    # Leer todos los datos y pasarlos a un DataFrame de Pandas
    data = sheet.get_all_values()
    df_sheet = pd.DataFrame(data[1:], columns=data[0])
    
    # 2. PROCESAMIENTO DE LA CONFIGURACIÓN (Lógica de columnas)
    if columna_seleccionada == 'produccion':
        aux_check = 7
    if columna_seleccionada == 'Geolocalizacion':
        aux_check = 8
    elif columna_seleccionada == 'Geomarketing':
        aux_check = 9

    df = pd.DataFrame({
        'capa': df_sheet.iloc[:, 0],
        'columna_db': df_sheet.iloc[:, 2],
        'check': df_sheet.iloc[:, aux_check]
    })

    df['capa'] = df['capa'].replace('', None).ffill()
    df['check'] = df['check'].astype(str).str.upper() == 'TRUE'

    # 3. ITERAR PARES (capa_bd, estructura_sheet, filtro_col, filtro_val)
    # Retorna lista de dicts: {'capa', 'estructura', 'filtro_col', 'filtro_val', 'ok', 'registros', 'error'}
    resultados = []

    for nombre_capa, estructura, filtro_col, filtro_val in pares:
        print(f"{nombre_capa.upper()} → estructura: {estructura}"
              + (f" | WHERE {filtro_col}='{filtro_val}'" if filtro_col else ""))

        # Filtrar columnas del Sheet usando la estructura elegida por el usuario
        columnas_validas = df[(df['capa'] == estructura) & (df['check'] == True)]

        if columnas_validas.empty:
            msg = (f"Sin columnas con TRUE para estructura='{estructura}' "
                   f"y columna='{columna_seleccionada}'")
            print(f"Capa {nombre_capa}: {msg}. Saltando...")
            resultados.append({'capa': nombre_capa, 'estructura': estructura,
                                'filtro_col': filtro_col, 'filtro_val': filtro_val,
                                'ok': False, 'registros': 0, 'error': msg})
            continue

        nombres_cols = ", ".join(columnas_validas['columna_db'].tolist())

        # Construir query con WHERE opcional
        query = f"SELECT {nombres_cols}, geom FROM {nombre_capa}"
        if filtro_col and filtro_val:
            query += f" WHERE {filtro_col} = '{filtro_val}'"

        #print(f"Leyendo [{nombre_capa}] con columnas de '{estructura}': {nombres_cols}...")
        #print(f"Esta es la QUERY: {query}")

        try:
            gdf = gpd.read_postgis(query, engine, geom_col='geom', crs='EPSG:4326')
            exportar_gdf(bd, nombre_capa, gdf, formato)
            print(f"Éxito: {len(gdf)} registros cargados.")
            resultados.append({'capa': nombre_capa, 'estructura': estructura,
                                'filtro_col': filtro_col, 'filtro_val': filtro_val,
                                'ok': True, 'registros': len(gdf), 'error': None})
        except Exception as e:
            print(f"Error al leer {nombre_capa}: {e}")
            error_msg = str(e).split('\n')[0][:120]
            resultados.append({'capa': nombre_capa, 'estructura': estructura,
                                'filtro_col': filtro_col, 'filtro_val': filtro_val,
                                'ok': False, 'registros': 0, 'error': error_msg})

    return resultados

# XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX
# XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX
# XXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX

# -> Aqui comienzan las funciones para completar el acta de entrega

import geopandas as gpd
import os
def procesar_carpeta(ruta_carpeta):
    # Extraer país del nombre de la carpeta (asumiendo: 'Cartografia_Colombia_Fecha')
    nombre_carpeta = os.path.basename(ruta_carpeta)
    pais = nombre_carpeta.split('_')[1]
    resumen_capas = []
    for archivo in os.listdir(ruta_carpeta):
        if archivo.endswith('.shp') or archivo.endswith('.gpkg') or archivo.endswith('.csv'):
            gdf = gpd.read_file(os.path.join(ruta_carpeta, archivo))
            conteo = len(gdf)
            resumen_capas.append(f"{archivo}: {conteo} registros")
    return pais, resumen_capas

import google.generativeai as genai

def generar_descripcion_ia(pais, resumen_capas, api_key):
    """
    Genera un párrafo técnico profesional sobre cartografía de Latinoamérica.
    """
    # --- SOLUCIÓN PARA EL ERROR DE CERTIFICADO EN REDES CORPORATIVAS ---
    # Fuerza a las librerías de Google a buscar los certificados del sistema (Windows)
    os.environ["REQUESTS_CA_BUNDLE"] = "sys"
    os.environ["GRPC_DEFAULT_SSL_ROOTS_FILE_PATH"] = "sys"
    
    # Configuración del modelo
    genai.configure(api_key=api_key)
    
    # Usamos 1.5 Flash por su baja latencia y excelente comprensión de contexto
    model = genai.GenerativeModel('gemini-2.5-flash')
    
    # Convertimos la lista de capas en un string legible para el prompt
    detalle_capas = ", ".join(resumen_capas)
    
    # Prompt de ingeniería: definimos rol, contexto y restricciones
    prompt = f"""
    Eres un Especialista en Sistemas de Información Geográfica (SIG) experto en cartografía de Latinoamérica.
    
    CONTEXTO:
    Se ha procesado una carpeta de archivos vectoriales (shapefiles) correspondiente a: {pais}.
    Las capas detectadas y su conteo de registros son: {detalle_capas}.
    
    TAREA:
    Escribe un párrafo técnico de máximo 4 líneas para un informe oficial. 
    Debe explicar la importancia de este conjunto de datos para el análisis territorial de {pais}. 
    Menciona de forma fluida algunas de las capas encontradas. 
    Usa un tono formal, profesional y preciso. No saludes, ve directo al grano.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"Error al generar descripción: {str(e)}"

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime
import os

def llenar_plantilla(ruta_plantilla, pais, resumen_capas, descripcion_gemini, ruta_carpeta, 
                    num_acta, usuario_nombre, usuario_correo):
    
    doc = Document(ruta_plantilla)
    fecha_str = datetime.now().strftime("%d/%m/%Y")
    texto_conteos = "\n".join(resumen_capas)
    
    # Agregamos los nuevos datos al diccionario de reemplazos
    reemplazos = {
        "{{FECHA}}": fecha_str,
        "{{PAIS}}": pais,
        "{{CONTEOS}}": texto_conteos,
        "{{INFO_CARTOGRAFICA}}": descripcion_gemini,
        "{{NUM_ACTA}}": str(num_acta),
        "{{USUARIO}}": usuario_nombre,
        "{{CORREO}}": usuario_correo
    }

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in reemplazos.items():
                        if key in p.text:
                            nuevo_texto = p.text.replace(key, value)
                            p.text = ""
                            run = p.add_run(nuevo_texto)
                            
                            # Formato Arial 12 Innegociable
                            run.font.name = 'Arial'
                            run.font.size = Pt(12)
                            r = run._element.get_or_add_rPr()
                            rFonts = r.get_or_add_rFonts()
                            rFonts.set(qn('w:ascii'), 'Arial')
                            rFonts.set(qn('w:hAnsi'), 'Arial')

    nombre_archivo = f"01_ACTA_ENTREGA_{pais}_{num_acta}.docx"
    ruta_final_salida = os.path.join(ruta_carpeta, nombre_archivo)
    
    doc.save(ruta_final_salida)
    return ruta_final_salida